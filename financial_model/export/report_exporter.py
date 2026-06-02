"""Word 报告导出引擎 — AllResults → 7章财务效益分析报告

从编排器结果生成标准 Word 报告, 结构:
  1. 基础参数复核 — 投资/融资/运营参数
  2. 盈利能力分析 — IRR/NPV/回收期
  3. 偿债能力分析 — DSCR 序列/还款计划
  4. 财务生存能力分析 — 现金流量
  5. 资产负债分析 — 资产负债率趋势
  6. 敏感性分析 — (可选, 接入 Phase 5 SensitivityResult)
  7. 结论与建议

典型用法::

    from financial_model.engines.orchestrator import ModelOrchestrator

    results = ModelOrchestrator.from_excel_v17().run()
    path = export_report(results, "报告.docx", project_name="某抽蓄项目")
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from financial_model.engines.orchestrator import AllResults


# ══════════════════════════════════════════════════════════
# 公共 API
# ══════════════════════════════════════════════════════════


def export_report(
    results: AllResults,
    path: str | Path,
    project_name: str = "",
    sensitivity_data: list[dict[str, Any]] | None = None,
) -> Path:
    """生成财务效益分析 Word 报告

    Args:
        results: 编排器的完整运行结果
        path: 输出 .docx 文件路径
        project_name: 项目名称
        sensitivity_data: 可选的敏感性分析数据 (Phase 5 接入)
            格式: [{"param": "上网电价", "negative": -0.01, "positive": 0.01}, ...]

    Returns:
        输出文件的 Path 对象
    """
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # 全局样式
    _setup_styles(doc)

    # 封面
    _write_cover(doc, project_name)

    # 第1-7章
    _write_chapter_1(doc, results)  # 基础参数复核
    _write_chapter_2(doc, results)  # 盈利能力分析
    _write_chapter_3(doc, results)  # 偿债能力分析
    _write_chapter_4(doc, results)  # 财务生存能力分析
    _write_chapter_5(doc, results)  # 资产负债分析
    _write_chapter_6(doc, sensitivity_data)  # 敏感性分析
    _write_chapter_7(doc, results)  # 结论

    doc.save(str(path))
    return path


# ══════════════════════════════════════════════════════════
# 样式
# ══════════════════════════════════════════════════════════


def _setup_styles(doc: Document) -> None:
    """配置文档默认样式"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)


# ══════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════


def _write_cover(doc: Document, project_name: str) -> None:
    """写入封面页"""
    # 标题
    title = doc.add_heading("财务效益分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 项目名称
    if project_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(project_name)
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()  # 空行


# ══════════════════════════════════════════════════════════
# 第1章: 基础参数复核
# ══════════════════════════════════════════════════════════


def _write_chapter_1(doc: Document, results: AllResults) -> None:
    doc.add_heading("第一章 基础参数复核", level=1)

    # 投资概算
    doc.add_heading("1.1 投资概算", level=2)
    invest_total = float(results.investment["construction_investment"].sum())
    fin = results.financing

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    _fill_table(table, [
        ("建设投资(万元)", f"{invest_total:,.2f}"),
        ("建设期利息(万元)", f"{fin.construction_interest_total:,.2f}"),
        ("动态总投资(万元)", f"{fin.dynamic_total_investment:,.2f}"),
        ("项目年限", f"{results.derived_metrics.project_years}年"),
    ])

    # 融资参数
    doc.add_heading("1.2 融资参数", level=2)
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = "Table Grid"
    _fill_table(table2, [
        ("建设期利率", f"{results.financing.construction_interest_total / fin.dynamic_total_investment * 100:.2f}% (利息/动态投资)"),
        ("还款期限", f"{len(results.financing.loan_schedule)}年"),
        ("还款方式", "等额本息"),
    ])


# ══════════════════════════════════════════════════════════
# 第2章: 盈利能力分析
# ══════════════════════════════════════════════════════════


def _write_chapter_2(doc: Document, results: AllResults) -> None:
    doc.add_heading("第二章 盈利能力分析", level=1)
    dm = results.derived_metrics

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    _fill_table(table, [
        ("全投资IRR", _fmt_pct(dm.irr_total)),
        ("资本金IRR", _fmt_pct(dm.irr_equity)),
        ("全投资NPV(万元)", _fmt_num(dm.npv_total)),
        ("资本金NPV(万元)", _fmt_num(dm.npv_equity)),
        ("静态回收期(年)", _fmt_years(dm.payback_static)),
        ("动态回收期(年)", _fmt_years(dm.payback_dynamic)),
    ])

    # 分析文字
    doc.add_paragraph()
    if dm.irr_total is not None:
        if dm.irr_total > dm.discount_rate:
            doc.add_paragraph(
                f"项目全投资IRR({dm.irr_total:.2%})高于基准收益率"
                f"({dm.discount_rate:.2%}), 项目在财务上可行。"
            )
        else:
            doc.add_paragraph(
                f"项目全投资IRR({dm.irr_total:.2%})低于基准收益率"
                f"({dm.discount_rate:.2%}), 项目在财务上不可行。"
            )


# ══════════════════════════════════════════════════════════
# 第3章: 偿债能力分析
# ══════════════════════════════════════════════════════════


def _write_chapter_3(doc: Document, results: AllResults) -> None:
    doc.add_heading("第三章 偿债能力分析", level=1)
    dm = results.derived_metrics

    # DSCR 摘要
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    _fill_table(table, [
        ("最低DSCR", _fmt_ratio(dm.dscr_min)),
        ("平均DSCR", _fmt_ratio(dm.dscr_avg)),
        ("还款期(年)", str(len(results.financing.loan_schedule))),
    ])

    doc.add_paragraph()
    if dm.dscr_min is not None:
        if dm.dscr_min >= 1.0:
            doc.add_paragraph(
                f"项目最低偿债备付率({dm.dscr_min:.2f})大于1.0, "
                "项目具备偿债能力。"
            )
        else:
            doc.add_paragraph(
                f"⚠ 项目最低偿债备付率({dm.dscr_min:.2f})低于1.0, "
                "存在偿债风险。"
            )

    # DSCR 年度表
    if dm.dscr_by_year:
        doc.add_heading("3.1 DSCR年度序列", level=2)
        dscr_items = sorted(dm.dscr_by_year.items())
        # 每10年一组
        dscr_table = doc.add_table(rows=len(dscr_items) + 1, cols=2)
        dscr_table.style = "Table Grid"
        dscr_table.rows[0].cells[0].text = "年度"
        dscr_table.rows[0].cells[1].text = "DSCR"
        for i, (year, ratio) in enumerate(dscr_items, 1):
            dscr_table.rows[i].cells[0].text = str(year)
            dscr_table.rows[i].cells[1].text = f"{ratio:.4f}"


# ══════════════════════════════════════════════════════════
# 第4章: 财务生存能力分析
# ══════════════════════════════════════════════════════════


def _write_chapter_4(doc: Document, results: AllResults) -> None:
    doc.add_heading("第四章 财务生存能力分析", level=1)

    cf = results.cf_plan.data

    # 检查累计盈余是否始终为正
    cumulative = cf["cumulative_surplus"]
    min_surplus = float(cumulative.min())
    max_surplus = float(cumulative.max())
    final_surplus = float(cumulative.iloc[-1])

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    _fill_table(table, [
        ("累计盈余最小值(万元)", f"{min_surplus:,.2f}"),
        ("累计盈余最大值(万元)", f"{max_surplus:,.2f}"),
        ("期末累计盈余(万元)", f"{final_surplus:,.2f}"),
    ])

    doc.add_paragraph()
    if min_surplus >= 0:
        doc.add_paragraph("项目财务计划现金流量各年累计盈余均非负, 财务生存能力良好。")
    else:
        doc.add_paragraph(
            f"⚠ 项目存在累计盈余为负的年度(最低{min_surplus:,.2f}万元), "
            "需关注短期融资安排。"
        )


# ══════════════════════════════════════════════════════════
# 第5章: 资产负债分析
# ══════════════════════════════════════════════════════════


def _write_chapter_5(doc: Document, results: AllResults) -> None:
    doc.add_heading("第五章 资产负债分析", level=1)
    dm = results.derived_metrics

    if dm.asset_liability_ratio:
        ratios = dm.asset_liability_ratio
        max_year = max(ratios, key=ratios.get)
        min_year = min(ratios, key=ratios.get)

        table = doc.add_table(rows=3, cols=3)
        table.style = "Table Grid"
        _fill_table_rows(table, [
            ["指标", "年度", "数值"],
            ["最高资产负债率", str(max_year), f"{ratios[max_year]:.2%}"],
            ["最低资产负债率", str(min_year), f"{ratios[min_year]:.2%}"],
        ])

        doc.add_paragraph()
        max_ratio = ratios[max_year]
        if max_ratio > 0.7:
            doc.add_paragraph(
                f"⚠ 最高资产负债率({max_ratio:.2%})较高, "
                "建议优化融资结构。"
            )
        else:
            doc.add_paragraph("项目资产负债率处于合理范围。")


# ══════════════════════════════════════════════════════════
# 第6章: 敏感性分析
# ══════════════════════════════════════════════════════════


def _write_chapter_6(
    doc: Document,
    sensitivity_data: list[dict[str, Any]] | None,
) -> None:
    doc.add_heading("第六章 敏感性分析", level=1)

    if not sensitivity_data:
        doc.add_paragraph("（未执行敏感性分析, 本章无数据。）")
        return

    # 敏感性参数影响表
    doc.add_heading("6.1 参数影响排序", level=2)
    doc.add_paragraph("以下为各参数对全投资IRR的影响幅度(龙卷风图数据):")

    table = doc.add_table(rows=len(sensitivity_data) + 1, cols=4)
    table.style = "Table Grid"
    _fill_table_rows(table, [
        ["参数", "负面影响", "正面影响", "影响幅度"],
        *[
            [
                item.get("param", ""),
                f"{item.get('negative', 0):+.4f}",
                f"{item.get('positive', 0):+.4f}",
                f"{item.get('spread', 0):.4f}",
            ]
            for item in sensitivity_data
        ],
    ])


# ══════════════════════════════════════════════════════════
# 第7章: 结论与建议
# ══════════════════════════════════════════════════════════


def _write_chapter_7(doc: Document, results: AllResults) -> None:
    doc.add_heading("第七章 结论与建议", level=1)
    dm = results.derived_metrics

    # 综合评价
    conclusions: list[str] = []

    if dm.irr_total is not None:
        if dm.irr_total > dm.discount_rate:
            conclusions.append(
                f"✅ 项目全投资IRR({dm.irr_total:.2%})高于基准收益率"
                f"({dm.discount_rate:.2%}), 盈利能力达标。"
            )
        else:
            conclusions.append(
                f"❌ 项目全投资IRR({dm.irr_total:.2%})低于基准收益率"
                f"({dm.discount_rate:.2%}), 盈利能力不足。"
            )

    if dm.dscr_min is not None:
        if dm.dscr_min >= 1.0:
            conclusions.append(
                f"✅ 最低DSCR({dm.dscr_min:.2f})≥1.0, 偿债能力达标。"
            )
        else:
            conclusions.append(
                f"❌ 最低DSCR({dm.dscr_min:.2f})<1.0, 存在偿债风险。"
            )

    if dm.payback_static is not None:
        conclusions.append(
            f"📊 静态投资回收期为{dm.payback_static:.1f}年。"
        )

    for c in conclusions:
        doc.add_paragraph(c, style="List Bullet")

    # 建议
    doc.add_heading("建议", level=2)
    suggestions: list[str] = []

    if dm.irr_total is not None and dm.irr_total < dm.discount_rate:
        suggestions.append("建议优化上网电价或降低建设投资以提高项目收益率。")
    if dm.dscr_min is not None and dm.dscr_min < 1.2:
        suggestions.append("建议适当延长还款期限或降低贷款利率以改善偿债能力。")
    if not suggestions:
        suggestions.append("项目财务指标整体良好, 建议按计划推进。")

    for s in suggestions:
        doc.add_paragraph(s, style="List Bullet")


# ══════════════════════════════════════════════════════════
# 辅助函数
# ══════════════════════════════════════════════════════════


def _fill_table(table: Any, rows: list[tuple[str, str]]) -> None:
    """填充两列表格 (label, value)"""
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value


def _fill_table_rows(table: Any, rows: list[list[str]]) -> None:
    """填充多列表格"""
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx].cells[c_idx].text = val


def _fmt_pct(v: float | None) -> str:
    if v is None:
        return "N/A"
    return f"{v:.2%}"


def _fmt_num(v: float) -> str:
    return f"{v:,.2f}"


def _fmt_ratio(v: float | None) -> str:
    if v is None:
        return "N/A"
    return f"{v:.2f}"


def _fmt_years(v: float | None) -> str:
    if v is None:
        return "N/A"
    return f"{v:.1f}"
