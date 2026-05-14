"""Export financial benefit analysis report as Word document.

Generates a structured .docx report matching the 财务效益分析.docx template:
1. 基础参数复核 (investment, financing, revenue, cost, tax parameters)
2. 盈利能力分析 (IRR, NPV, payback period)
3. 偿债能力分析 (DSCR series, loan repayment)
4. 财务生存能力分析 (cash flow analysis)
5. 敏感性分析 (parameter perturbation impact table)
"""
from __future__ import annotations

import os
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from financial_kg.models.graph import FinancialGraph
from financial_kg.engine.derived_metrics import (
    compute_derived_metrics,
    serialize_metrics,
    deserialize_metrics,
    DerivedMetrics,
)
from financial_kg.engine.sensitivity import (
    run_sensitivity,
    SensitivityResult,
    _build_spider_table,
)


# ── Report generation ────────────────────────────────────────────────────────

def export_financial_report(
    graph: FinancialGraph,
    output_path: str,
    task_id: str = "",
    output_dir: str = "",
    snapshots_dir: str = "",
    project_name: str = "",
    sensitivity_params: list[tuple[str, str]] | None = None,
    sensitivity_result: SensitivityResult | None = None,
) -> str:
    """Generate a financial benefit analysis Word report.

    Args:
        graph: The FinancialGraph with current values.
        output_path: Path for the output .docx file.
        task_id: Task ID for sensitivity snapshot creation.
        output_dir: Output directory.
        snapshots_dir: Snapshots directory root.
        project_name: Project name for the report header.
        sensitivity_params: List of (cell_id, display_name) for sensitivity analysis.
            Defaults to common parameters if not specified.
        sensitivity_result: Pre-computed sensitivity result to skip recalculation.

    Returns:
        Path to the generated .docx file.
    """
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    metrics = compute_derived_metrics(graph)

    # ── Title ─────────────────────────────────────────────────────────
    title = doc.add_heading("财务效益分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if project_name:
        doc.add_paragraph(f"项目名称：{project_name}").runs[0].font.size = Pt(12)

    # ── 1. 基础参数复核 ──────────────────────────────────────────────
    _add_section(doc, "一、基础参数复核", level=1)

    param_sections = [
        ("投资参数", _INVESTMENT_KEYWORDS),
        ("融资参数", _FINANCE_KEYWORDS),
        ("收入参数", _REVENUE_KEYWORDS),
        ("成本参数", _COST_KEYWORDS),
        ("税收参数", _TAX_KEYWORDS),
    ]

    for section_name, keywords in param_sections:
        _add_param_table(doc, graph, section_name, keywords)

    # ── 2. 盈利能力分析 ──────────────────────────────────────────────
    _add_section(doc, "二、盈利能力分析", level=1)
    _add_profitability_table(doc, metrics)

    # ── 3. 偿债能力分析 ──────────────────────────────────────────────
    _add_section(doc, "三、偿债能力分析", level=1)
    _add_dscr_table(doc, metrics)

    # ── 4. 财务生存能力分析 ──────────────────────────────────────────
    _add_section(doc, "四、财务生存能力分析", level=1)
    _add_cashflow_table(doc, metrics)

    # ── 5. 敏感性分析 ────────────────────────────────────────────────
    _add_section(doc, "五、敏感性分析", level=1)

    if sensitivity_params or sensitivity_result:
        if sensitivity_result:
            _add_sensitivity_from_result(doc, metrics, sensitivity_result)
        else:
            params = sensitivity_params
            if params is None:
                params = _auto_detect_params(graph)
            if params:
                _add_sensitivity_section(doc, graph, metrics, params, task_id, output_dir, snapshots_dir)
            else:
                p = doc.add_paragraph("未检测到可分析的输入参数。")
                p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    else:
        # No params/result provided — auto-detect and run
        params = _auto_detect_params(graph)
        if params:
            _add_sensitivity_section(doc, graph, metrics, params, task_id, output_dir, snapshots_dir)
        else:
            p = doc.add_paragraph("未检测到可分析的输入参数。")
            p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # Save
    doc.save(output_path)
    return output_path


# ── Keyword lists ────────────────────────────────────────────────────────────

_INVESTMENT_KEYWORDS = [
    "静态总投资", "动态总投资", "EPC合同额", "建设投资", "工程总投资",
]
_FINANCE_KEYWORDS = [
    "贷款利率", "还款期限", "宽限期", "资本金比例", "融资比例",
]
_REVENUE_KEYWORDS = [
    "售电电价", "售电电量", "加权平均电价", "营业收入", "电费收入",
]
_COST_KEYWORDS = [
    "综合购电成本", "维修费", "折旧费", "人工成本", "经营成本",
]
_TAX_KEYWORDS = [
    "所得税率", "增值税率", "税金及附加", "所得税",
]


# ── Section builders ─────────────────────────────────────────────────────────

def _add_section(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_param_table(
    doc: Document,
    graph: FinancialGraph,
    section_name: str,
    keywords: list[str],
) -> None:
    """Add a parameter verification table for one section."""
    matches = []
    for ind in graph.indicators.values():
        name = ind.name or ""
        if any(kw in name for kw in keywords):
            val = ind.summary_value
            if val is not None:
                matches.append((name, val, ind.unit or ""))

    if not matches:
        return

    table = doc.add_table(rows=len(matches) + 1, cols=3, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ["参数名称", "数值", "单位"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    for row_idx, (name, val, unit) in enumerate(matches, 1):
        table.rows[row_idx].cells[0].text = name[:30]
        if isinstance(val, float):
            table.rows[row_idx].cells[1].text = f"{val:,.2f}"
        else:
            table.rows[row_idx].cells[1].text = str(val)
        table.rows[row_idx].cells[2].text = unit


def _add_profitability_table(doc: Document, metrics: DerivedMetrics) -> None:
    """Add profitability analysis table."""
    rows_data = [
        ("税后全投资内部收益率（IRR）", metrics.irr_after_tax, "%"),
        ("财务净现值（NPV）", metrics.npv_after_tax, ""),
        ("投资回收期", metrics.payback_period, "年"),
        ("年均净现金流", metrics.annual_net_cashflow, ""),
    ]

    table = doc.add_table(rows=len(rows_data) + 1, cols=3, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["指标", "数值", "单位"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    for row_idx, (name, val, unit) in enumerate(rows_data, 1):
        table.rows[row_idx].cells[0].text = name
        if val is not None:
            if unit == "%":
                table.rows[row_idx].cells[1].text = f"{val * 100:.2f}"
            elif isinstance(val, float):
                table.rows[row_idx].cells[1].text = f"{val:,.2f}"
            else:
                table.rows[row_idx].cells[1].text = str(val)
        else:
            table.rows[row_idx].cells[1].text = "—"
        table.rows[row_idx].cells[2].text = unit


def _add_dscr_table(doc: Document, metrics: DerivedMetrics) -> None:
    """Add DSCR analysis table."""
    rows_data = [
        ("DSCR 均值", metrics.dscr_avg, ""),
        ("DSCR 最低值", metrics.dscr_min, ""),
    ]

    table = doc.add_table(rows=len(rows_data) + 1, cols=3, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["指标", "数值", "单位"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    for row_idx, (name, val, _unit) in enumerate(rows_data, 1):
        table.rows[row_idx].cells[0].text = name
        if val is not None:
            table.rows[row_idx].cells[1].text = f"{val:.2f}"
        else:
            table.rows[row_idx].cells[1].text = "—"
        table.rows[row_idx].cells[2].text = ""

    # Annual DSCR series
    if metrics.dscr_series:
        doc.add_paragraph("")
        p = doc.add_paragraph("DSCR 年度分布：")
        p.runs[0].bold = True

        annual = sorted(metrics.dscr_series.items())
        table = doc.add_table(rows=2, cols=len(annual) + 1, style="Light Grid Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        table.rows[0].cells[0].text = "指标"
        for run in table.rows[0].cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

        for col_idx, (year, _) in enumerate(annual, 1):
            cell = table.rows[0].cells[col_idx]
            cell.text = str(year)
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)

        table.rows[1].cells[0].text = "DSCR"
        for col_idx, (_, val) in enumerate(annual, 1):
            table.rows[1].cells[col_idx].text = f"{val:.2f}"


def _add_cashflow_table(doc: Document, metrics: DerivedMetrics) -> None:
    """Add cash flow survival analysis table."""
    rows_data = [
        ("动态总投资", metrics.total_investment_dynamic, ""),
        ("静态总投资", metrics.total_investment_static, ""),
        ("全期营业收入", metrics.total_revenue, ""),
        ("全期总成本", metrics.total_cost, ""),
        ("全期税费", metrics.total_tax, ""),
        ("年均净现金流", metrics.annual_net_cashflow, ""),
    ]

    table = doc.add_table(rows=len(rows_data) + 1, cols=3, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["指标", "数值", "单位"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    for row_idx, (name, val, _unit) in enumerate(rows_data, 1):
        table.rows[row_idx].cells[0].text = name
        if val is not None:
            table.rows[row_idx].cells[1].text = f"{val:,.2f}"
        else:
            table.rows[row_idx].cells[1].text = "—"
        table.rows[row_idx].cells[2].text = ""


def _add_sensitivity_section(
    doc: Document,
    graph: FinancialGraph,
    base_metrics: DerivedMetrics,
    param_cells: list[tuple[str, str]],
    task_id: str = "",
    output_dir: str = "",
    snapshots_dir: str = "",
) -> None:
    """Run sensitivity analysis and add results to report."""
    result = run_sensitivity(
        graph=graph,
        param_cells=param_cells,
        perturbations=[-0.1, -0.05, 0.05, 0.1],
        task_id=task_id,
        output_dir=output_dir,
        snapshots_dir=snapshots_dir,
    )

    doc.add_paragraph("")
    p = doc.add_paragraph("敏感性分析结果（IRR）：")
    p.runs[0].bold = True

    spider = _build_spider_table(base_metrics, result.scenarios, "irr_after_tax")
    if spider:
        cols_count = len(spider[0]) if spider else 0
        table = doc.add_table(rows=len(spider) + 1, cols=cols_count, style="Light Grid Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        headers = list(spider[0].keys())
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)

        for row_idx, row_data in enumerate(spider, 1):
            for col_idx, h in enumerate(headers):
                val = row_data.get(h, "")
                table.rows[row_idx].cells[col_idx].text = str(val) if val is not None else "—"
                for run in table.rows[row_idx].cells[col_idx].paragraphs[0].runs:
                    run.font.size = Pt(9)


def _add_sensitivity_from_result(
    doc: Document,
    base_metrics: DerivedMetrics,
    result: SensitivityResult,
) -> None:
    """Add pre-computed sensitivity results to report (no recalculation)."""
    doc.add_paragraph("")
    p = doc.add_paragraph("敏感性分析结果（IRR）：")
    p.runs[0].bold = True

    spider = _build_spider_table(base_metrics, result.scenarios, "irr_after_tax")
    if spider:
        cols_count = len(spider[0]) if spider else 0
        table = doc.add_table(rows=len(spider) + 1, cols=cols_count, style="Light Grid Accent 1")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = list(spider[0].keys())
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)

        for row_idx, row_data in enumerate(spider, 1):
            for col_idx, h in enumerate(headers):
                val = row_data.get(h, "")
                table.rows[row_idx].cells[col_idx].text = str(val) if val is not None else "—"
                for run in table.rows[row_idx].cells[col_idx].paragraphs[0].runs:
                    run.font.size = Pt(9)


# ── Auto-detect parameters for sensitivity analysis ──────────────────────────

def _auto_detect_params(graph: FinancialGraph) -> list[tuple[str, str]] | None:
    """Auto-detect key input parameters for sensitivity analysis."""
    keywords = [
        "电价", "售电电价", "上网电价",
        "电量", "售电电量", "发电量",
        "贷款利率", "利率",
        "成本", "购电成本", "经营成本",
        "静态投资", "动态投资", "总投资",
    ]

    results: list[tuple[str, str]] = []
    seen_params: set[str] = set()

    for ind in graph.indicators.values():
        name = ind.name or ""
        if not any(kw in name for kw in keywords):
            continue
        if name in seen_params:
            continue

        # Find first numeric non-formula cell in this indicator
        for cid in ind.cell_ids:
            cell = graph.cells.get(cid)
            if cell and not cell.formula_raw and isinstance(cell.value, (int, float)) and cell.value != 0:
                results.append((cid, name))
                seen_params.add(name)
                break

    return results if results else None
