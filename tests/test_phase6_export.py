"""Phase 6 测试: 导出层 — excel_exporter + report_exporter

验证:
  1. Excel: 文件生成, sheet数量/名称, 行列数, 派生指标值一致
  2. Word: 文件生成, 7章标题, 关键指标数值, 表格存在
  3. 集成: 完整管道 (params → run → export → read → verify)
"""
from __future__ import annotations

import os
import tempfile

import pytest

from financial_model.engines.orchestrator import AllResults, ModelOrchestrator
from financial_model.export.excel_exporter import export_excel
from financial_model.export.report_exporter import export_report


# ══════════════════════════════════════════════════════════
# Fixtures
# ══════════════════════════════════════════════════════════


@pytest.fixture
def results() -> AllResults:
    """黄金基准结果"""
    return ModelOrchestrator.from_excel_v17().run()


@pytest.fixture
def tmp_dir() -> str:
    """临时目录"""
    import shutil
    d = tempfile.mkdtemp()
    yield d
    # 清理 (递归, 处理子目录)
    shutil.rmtree(d, ignore_errors=True)


# ══════════════════════════════════════════════════════════
# Excel Export Tests
# ══════════════════════════════════════════════════════════


class TestExcelExport:
    """export_excel 基本功能"""

    def test_creates_file(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        path = os.path.join(tmp_dir, "test.xlsx")
        result_path = export_excel(results, path)
        assert os.path.exists(result_path)

    def test_file_not_empty(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        path = os.path.join(tmp_dir, "test.xlsx")
        export_excel(results, path)
        assert os.path.getsize(path) > 1000  # 至少 1KB

    def test_sheet_count(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """应有 13 个 sheet"""
        from openpyxl import load_workbook

        path = os.path.join(tmp_dir, "test.xlsx")
        export_excel(results, path)
        wb = load_workbook(path)
        assert len(wb.sheetnames) == 13

    def test_sheet_names(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """关键 sheet 应存在"""
        from openpyxl import load_workbook

        path = os.path.join(tmp_dir, "test.xlsx")
        export_excel(results, path)
        wb = load_workbook(path)
        names = wb.sheetnames
        # 第一个 sheet = 摘要
        assert "项目摘要" in names or names[0].endswith("摘要")
        # 最后一个 = 派生指标
        assert "派生指标" in names[-1] or names[-1].endswith("指标")

    def test_data_sheet_row_counts(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """数据表应有正确行数 (标题3行 + 48数据行 = 51行)"""
        from openpyxl import load_workbook

        path = os.path.join(tmp_dir, "test.xlsx")
        export_excel(results, path)
        wb = load_workbook(path)

        # 折旧表 (48行数据)
        depr_sheet = None
        for name in wb.sheetnames:
            if "折旧" in name:
                depr_sheet = wb[name]
                break
        assert depr_sheet is not None
        # 3行(title+header+blank) + 48 data = 51
        assert depr_sheet.max_row >= 51

    def test_derived_metrics_values(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """派生指标 sheet 中的值应与 AllResults 一致"""
        from openpyxl import load_workbook

        path = os.path.join(tmp_dir, "test.xlsx")
        export_excel(results, path)
        wb = load_workbook(path)

        # 找到派生指标 sheet
        dm_sheet = wb[wb.sheetnames[-1]]
        # 读取所有单元格文本
        texts = []
        for row in dm_sheet.iter_rows(values_only=True):
            texts.extend([str(v) for v in row if v is not None])

        # 验证 IRR 值出现
        dm = results.derived_metrics
        if dm.irr_total is not None:
            irr_str = f"{dm.irr_total:.2%}"
            assert irr_str in texts, f"IRR {irr_str} not found in sheet"

    def test_project_name_in_summary(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """项目名称应出现在摘要 sheet"""
        from openpyxl import load_workbook

        path = os.path.join(tmp_dir, "test.xlsx")
        export_excel(results, path, project_name="测试抽蓄项目")
        wb = load_workbook(path)
        summary = wb[wb.sheetnames[0]]
        texts = []
        for row in summary.iter_rows(values_only=True):
            texts.extend([str(v) for v in row if v is not None])
        assert any("测试抽蓄项目" in t for t in texts)

    def test_creates_parent_directory(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """应自动创建不存在的父目录"""
        nested_path = os.path.join(tmp_dir, "sub", "dir", "test.xlsx")
        result_path = export_excel(results, nested_path)
        assert os.path.exists(result_path)


class TestExcelDataIntegrity:
    """Excel 数据完整性验证"""

    def test_investment_sheet_has_data(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """投资概算表应有数据"""
        from openpyxl import load_workbook

        path = os.path.join(tmp_dir, "test.xlsx")
        export_excel(results, path)
        wb = load_workbook(path)

        inv_sheet = None
        for name in wb.sheetnames:
            if "投资" in name:
                inv_sheet = wb[name]
                break
        assert inv_sheet is not None
        # 应有超过3行 (标题+表头+数据)
        assert inv_sheet.max_row > 3

    def test_balance_sheet_has_17_columns(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """资产负债表应有 17 列数据 + 1 列年份 = 18列"""
        from openpyxl import load_workbook

        path = os.path.join(tmp_dir, "test.xlsx")
        export_excel(results, path)
        wb = load_workbook(path)

        bs_sheet = None
        for name in wb.sheetnames:
            if "资产负债" in name:
                bs_sheet = wb[name]
                break
        assert bs_sheet is not None
        # max_column >= 18 (year + 17 data cols)
        assert bs_sheet.max_column >= 18


# ══════════════════════════════════════════════════════════
# Word Report Tests
# ══════════════════════════════════════════════════════════


class TestReportExport:
    """export_report 基本功能"""

    def test_creates_file(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        path = os.path.join(tmp_dir, "test.docx")
        result_path = export_report(results, path)
        assert os.path.exists(result_path)

    def test_file_not_empty(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        path = os.path.join(tmp_dir, "test.docx")
        export_report(results, path)
        assert os.path.getsize(path) > 1000

    def test_has_seven_chapters(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """应有 7 章标题"""
        from docx import Document

        path = os.path.join(tmp_dir, "test.docx")
        export_report(results, path)
        doc = Document(path)
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        # 7 章 + 子标题
        chapter_headings = [h for h in headings if h.startswith("第")]
        assert len(chapter_headings) == 7

    def test_chapter_titles(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """7章标题应包含关键词"""
        from docx import Document

        path = os.path.join(tmp_dir, "test.docx")
        export_report(results, path)
        doc = Document(path)
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        all_text = " ".join(headings)
        # 验证关键章节
        assert "盈利" in all_text or "基础" in all_text
        assert "偿债" in all_text or "DSCR" in all_text
        assert "结论" in all_text

    def test_has_tables(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """应包含数据表格"""
        from docx import Document

        path = os.path.join(tmp_dir, "test.docx")
        export_report(results, path)
        doc = Document(path)
        assert len(doc.tables) >= 5  # 至少5个表格

    def test_irr_value_in_report(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """IRR 值应出现在报告中"""
        from docx import Document

        path = os.path.join(tmp_dir, "test.docx")
        export_report(results, path)
        doc = Document(path)
        all_text = " ".join(p.text for p in doc.paragraphs)
        dm = results.derived_metrics
        if dm.irr_total is not None:
            irr_str = f"{dm.irr_total:.2%}"
            assert irr_str in all_text, f"IRR {irr_str} not found in report"

    def test_dscr_value_in_report(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """DSCR 值应出现在报告中"""
        from docx import Document

        path = os.path.join(tmp_dir, "test.docx")
        export_report(results, path)
        doc = Document(path)
        all_text = " ".join(p.text for p in doc.paragraphs)
        dm = results.derived_metrics
        if dm.dscr_min is not None:
            dscr_str = f"{dm.dscr_min:.2f}"
            assert dscr_str in all_text, f"DSCR {dscr_str} not found in report"

    def test_project_name_in_cover(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """项目名称应出现在封面"""
        from docx import Document

        path = os.path.join(tmp_dir, "test.docx")
        export_report(results, path, project_name="某抽蓄电站")
        doc = Document(path)
        # 前几个段落应包含项目名
        early_text = " ".join(p.text for p in doc.paragraphs[:5])
        assert "某抽蓄电站" in early_text

    def test_sensitivity_chapter_with_data(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """第6章应显示敏感性数据"""
        from docx import Document

        path = os.path.join(tmp_dir, "test.docx")
        tornado_data = [
            {"param": "上网电价", "negative": -0.01, "positive": 0.01, "spread": 0.02},
        ]
        export_report(results, path, sensitivity_data=tornado_data)
        doc = Document(path)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "上网电价" in all_text

    def test_sensitivity_chapter_without_data(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """第6章无数据时应显示提示文字"""
        from docx import Document

        path = os.path.join(tmp_dir, "test.docx")
        export_report(results, path, sensitivity_data=None)
        doc = Document(path)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "未执行" in all_text

    def test_creates_parent_directory(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """应自动创建不存在的父目录"""
        nested_path = os.path.join(tmp_dir, "reports", "sub", "test.docx")
        result_path = export_report(results, nested_path)
        assert os.path.exists(result_path)


class TestReportContent:
    """报告内容验证"""

    def test_conclusion_section(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """结论章应有 IRR 判断"""
        from docx import Document

        path = os.path.join(tmp_dir, "test.docx")
        export_report(results, path)
        doc = Document(path)
        all_text = " ".join(p.text for p in doc.paragraphs)

        dm = results.derived_metrics
        if dm.irr_total is not None:
            if dm.irr_total > dm.discount_rate:
                assert "达标" in all_text or "高于" in all_text
            else:
                assert "不足" in all_text or "低于" in all_text

    def test_dscr_annual_table(
        self, results: AllResults, tmp_dir: str
    ) -> None:
        """DSCR 年度表应存在且有数据行"""
        from docx import Document

        path = os.path.join(tmp_dir, "test.docx")
        export_report(results, path)
        doc = Document(path)

        dm = results.derived_metrics
        if dm.dscr_by_year:
            # 应有至少一个包含 DSCR 数据的表格
            found_dscr_table = False
            for table in doc.tables:
                if any(
                    "DSCR" in cell.text
                    for row in table.rows
                    for cell in row.cells
                ):
                    found_dscr_table = True
                    break
            assert found_dscr_table
